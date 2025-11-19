import io
import logging
from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from PyPDF2 import PdfReader
from sqlalchemy.ext.asyncio import AsyncSession
from tqdm import tqdm

from db.db_models import Document, DocumentChunk, Session
from db.vectordb import vectordb
from services.llm_service import llm_service

logger = logging.getLogger(__name__)


class IngestionService:


    async def ingest(self, file, db: AsyncSession, session_id: str):

        logger.info(f"Starting ingestion for file: {file.filename}")

        #1. Extract raw text
        raw_text = await self._extract_text(file)
        logger.info(f"Extracted {len(raw_text)} characters from uploaded file")

        #2. Clean text
        cleaned = self._clean_text(raw_text)
        logger.info(f"Cleaned text length: {len(cleaned)}")

        #3. Chunk text
        chunks = self._chunk_text(cleaned)
        logger.info(f"Generated {len(chunks)} chunks")

        #4. Adding Document to SQLite DB
        doc = Document(filename=file.filename, 
                       content_type=file.content_type,
                       session_id = session_id)
        db.add(doc)
        await db.flush() #added to fetch auto-incremented ID in next step
        doc_id = doc.id
        logger.info(f"Added document to Document table: {doc_id}")

        
        #5. Save chunks + embeddings into ChromaDB
        logger.info(f"Adding all DocumentChunk objects (text) to SQLite DB")
        logger.info(f"Creating chunk embeddings and adding them to ChromaDB")
        for idx, chunk_text in enumerate(tqdm(chunks, desc="Embedding chunks", unit="chunk")):
            
            #Adding DocumentChunk (text) to SQLite DB
            chunk = DocumentChunk(document_id=doc_id, session_id=session_id, 
                                  chunk_index=idx, text=chunk_text)
            db.add(chunk)
            await db.flush() #added to fetch auto-incremented ID in next step
            chunk_id = chunk.id

            emb = await llm_service.embed(chunk_text)

            metadata = {"session_id": session_id,
                        "doc_id": doc_id,
                        "chunk_id": chunk_id,
                        "chunk_index": idx,
                        "text": chunk_text}

            vector_id = f"{session_id}_{doc_id}_{chunk_id}"

            vectordb.add_vector(embedding=emb, collection_name='chunks', metadata=metadata, vector_id=vector_id)

        logger.info(f"Completed ChromaDB ingestion: {len(chunks)} vectors added")

        #Generate a name for the session/chat
        try:
            preview_text = chunks[0][:500]
            session_name = await llm_service.generate_session_title(preview_text)
            # Update session row with name
            session_row = await db.get(Session, session_id)
            session_row.session_name = session_name
            await db.commit()
            logger.info(f"[{session_id}] Generated session name: {session_name}")
        except Exception as e:
            logger.error(f"Failed to generate session name: {e}")

        await db.commit()
        logger.info(f"Inserted {len(chunks)} chunks into 'document_chunks' table for document: {doc_id}")
        logger.info(f"Finished ingestion for document: {doc_id}")

        return doc_id


    async def _extract_text(self, file):
        """
        Extract text from .pdf, .txt, .docx, .md files.
        """

        content = await file.read()
        filename = file.filename.lower()

        #Txt and Markdown  file formats
        if filename.endswith(".txt") or filename.endswith(".md"):
            logger.info("Detected TXT/MD file")
            return content.decode("utf-8", errors="ignore")

        #DOCX file format
        if filename.endswith(".docx"):
            logger.info("Detected DOCX file")
            doc = DocxDocument(io.BytesIO(content))
            return "\n".join(par.text for par in doc.paragraphs)

        #PDF file format
        if filename.endswith(".pdf"):
            logger.info("Detected PDF file")
            reader = PdfReader(io.BytesIO(content))
            extracted = []
            for page in reader.pages:
                text = page.extract_text() or ""
                extracted.append(text)
            return "\n".join(extracted)

        #Unsupported file format
        logger.warning("Unsupported file type")
        return ""



    def _clean_text(self, text):
        return text.replace("\r", "").strip()


    def _chunk_text(self, text):
        splitter = RecursiveCharacterTextSplitter(chunk_size=1000,
                                                  chunk_overlap=150,
                                                  length_function=len)
        return splitter.split_text(text)


ingestion_service = IngestionService()
